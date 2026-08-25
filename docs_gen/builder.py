# -*- coding: utf-8 -*-
"""Builder: يحول قوائم المحتوى إلى ملف Word عربي RTL شامل."""
import sys, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

AR_FONT = "Segoe UI"
BLUE = RGBColor(0x0E, 0x63, 0x8F)
DARK = RGBColor(0x0B, 0x19, 0x29)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
SUBTLE = RGBColor(0x55, 0x65, 0x74)

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    for run in paragraph.runs:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), AR_FONT)
        rFonts.set(qn('w:hAnsi'), AR_FONT)
        rFonts.set(qn('w:cs'), AR_FONT)
        # w:rtl على مستوى الـ run: يضمن عرض الأقواس واتجاه النص المختلط
        # (عربي + إنجليزي بين قوسين) بشكل صحيح في Word بدل ما تظهر معكوسة )(
        if rPr.find(qn('w:rtl')) is None:
            rPr.append(OxmlElement('w:rtl'))

def style_run(run, size, bold=False, color=None, italic=False):
    run.font.name = AR_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color

ALIGN_MAP = {
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

def to_color(val):
    if isinstance(val, tuple):
        return RGBColor(*val)
    return val

def add_paragraph(doc, text, size=11, bold=False, color=None, align=None, space_after=6, italic=False, indent=0.0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = ALIGN_MAP.get(align, align)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    style_run(r, size, bold, to_color(color), italic)
    set_rtl(p)
    return p

def add_bullet(doc, text, size=11, indent=0.6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run("•  " + text)
    style_run(r, size)
    set_rtl(p)
    return p

def add_step(doc, num, text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{num} -  ")
    style_run(r1, size, bold=True, color=BLUE)
    r2 = p.add_run(text)
    style_run(r2, size)
    set_rtl(p)
    return p

def add_heading(doc, text, level):
    sizes = {1: 20, 2: 16, 3: 13.5}
    colors = {1: BLUE, 2: DARK, 3: GOLD}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    style_run(r, sizes[level], bold=True, color=colors[level])
    set_rtl(p)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.3
    r1 = p.add_run("💡 ملاحظة: ")
    style_run(r1, 10.5, bold=True, color=GOLD)
    r2 = p.add_run(text)
    style_run(r2, 10.5, color=SUBTLE, italic=True)
    set_rtl(p)
    return p

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        style_run(r, 10.5, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_rtl(p)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), '0E638F')
        cell._tc.get_or_add_tcPr().append(shd)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            style_run(r, 10.5)
            set_rtl(p)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def build(all_sections, out_path):
    doc = Document()
    # إعداد الصفحة
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    for block in all_sections:
        kind = block.get("kind")
        if kind == "title":
            add_paragraph(doc, block["text"], 26, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
        elif kind == "subtitle":
            add_paragraph(doc, block["text"], 13, color=GOLD, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
        elif kind == "pagebreak":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        elif kind == "h1":
            add_heading(doc, block["title"], 1)
        elif kind == "h2":
            add_heading(doc, block["title"], 2)
        elif kind == "h3":
            add_heading(doc, block["title"], 3)
        elif kind == "p":
            add_paragraph(doc, block["text"], block.get("size", 11), block.get("bold", False),
                          block.get("color"), block.get("align"), block.get("space_after", 6))
        elif kind == "bullet":
            add_bullet(doc, block["text"], block.get("size", 11))
        elif kind == "steps":
            for i, s in enumerate(block["items"], 1):
                add_step(doc, i, s, block.get("size", 11))
        elif kind == "note":
            add_note(doc, block["text"])
        elif kind == "table":
            add_table(doc, block["headers"], block["rows"])
        elif kind == "spacer":
            add_paragraph(doc, "", size=6, space_after=2)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path
